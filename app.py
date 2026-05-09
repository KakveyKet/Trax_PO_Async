import streamlit as st
import pandas as pd
import re
from PyPDF2 import PdfReader
import docx
import io
import sqlite3
from datetime import datetime
import os
import openpyxl

# Set up the page layout with a professional icon
st.set_page_config(page_title="PO Auto-Sync System", page_icon=":material/account_tree:", layout="wide")

# --- DATABASE SETUP ---
DB_NAME = "po_processing_log.db"

def init_db():
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS processing_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT, file_name TEXT, processed_date TEXT, po_count INTEGER)''')
    
    cursor.execute('''CREATE TABLE IF NOT EXISTS extracted_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT, history_id INTEGER, file_name TEXT, po_number TEXT, size TEXT, quantity TEXT)''')
    
    cursor.execute("PRAGMA table_info(extracted_data)")
    columns = [col[1] for col in cursor.fetchall()]
    if 'size' not in columns:
        cursor.execute("ALTER TABLE extracted_data ADD COLUMN size TEXT DEFAULT ''")

    cursor.execute('''CREATE TABLE IF NOT EXISTS master_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT, file_name TEXT, upload_date TEXT, file_data BLOB)''')
    conn.commit()
    conn.close()

def log_to_db(filename, file_results_df):
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    po_count = len(file_results_df)
    
    cursor.execute("INSERT INTO processing_history (file_name, processed_date, po_count) VALUES (?, ?, ?)", (filename, current_time, po_count))
    history_id = cursor.lastrowid
    
    for index, row in file_results_df.iterrows():
        size_val = row.get('Size', '')
        if pd.isna(size_val): size_val = ""
        cursor.execute("INSERT INTO extracted_data (history_id, file_name, po_number, size, quantity) VALUES (?, ?, ?, ?, ?)", 
                       (history_id, filename, row['PO number'], size_val, row['QTY']))
    conn.commit()
    conn.close()

def fetch_db_report():
    conn = sqlite3.connect(DB_NAME)
    df = pd.read_sql_query("SELECT id, file_name as 'File Name', processed_date as 'Date Processed', po_count as 'PO Rows Extracted' FROM processing_history ORDER BY id DESC", conn)
    conn.close()
    return df

def fetch_specific_file_data(history_id):
    conn = sqlite3.connect(DB_NAME)
    df = pd.read_sql_query(f"SELECT file_name as 'File Name', po_number as 'PO number', size as 'Size', quantity as 'QTY' FROM extracted_data WHERE history_id = {history_id}", conn)
    conn.close()
    return df

def delete_history_record(history_id):
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM extracted_data WHERE history_id = ?", (history_id,))
    cursor.execute("DELETE FROM processing_history WHERE id = ?", (history_id,))
    conn.commit()
    conn.close()

def get_master_files_summary():
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        cursor.execute("SELECT count(name) FROM sqlite_master WHERE type='table' AND name='master_po_data'")
        if cursor.fetchone()[0] == 0: return pd.DataFrame() 
        query = "SELECT Source_File_Name as 'File Name', MAX(Upload_Date) as 'Upload Date', COUNT(*) as 'Total Rows' FROM master_po_data GROUP BY Source_File_Name ORDER BY MAX(Upload_Date) DESC"
        df = pd.read_sql_query(query, conn)
        conn.close()
        return df
    except: return pd.DataFrame()

def delete_master_file(filename):
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    try: cursor.execute("DELETE FROM master_po_data WHERE Source_File_Name = ?", (filename,))
    except: pass
    cursor.execute("DELETE FROM master_files WHERE file_name = ?", (filename,))
    conn.commit()
    conn.close()

def preview_master_file(filename):
    conn = sqlite3.connect(DB_NAME)
    df = pd.read_sql_query("SELECT * FROM master_po_data WHERE Source_File_Name = ?", conn, params=(filename,))
    conn.close()
    return df

def get_raw_master_file(filename):
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        cursor.execute("SELECT file_data FROM master_files WHERE file_name = ?", (filename,))
        row = cursor.fetchone()
        conn.close()
        return row[0] if row else None
    except: return None

# --- BULLETPROOF SIZE NORMALIZATION ENGINE ---
def normalize_size(size_str):
    """
    Safely converts formats like '45cm x 2.7cm' or '2.7X45CM' 
    into a perfect matching code: '2.7X45'
    """
    if not size_str or pd.isna(size_str) or str(size_str).strip() in ['None', '']: 
        return ""
    
    s = str(size_str).upper()
    
    # 1. Look specifically for the [Number] x [Number] pattern
    match = re.search(r'(\d+(?:\.\d+)?)\s*[a-zA-Z]*\s*[xX*]\s*(\d+(?:\.\d+)?)', s)
    if match:
        val1 = float(match.group(1))
        val2 = float(match.group(2))
        
        # Sort smallest to largest
        vals = sorted([val1, val2])
        
        # Format cleanly (remove .0 from whole numbers)
        clean_nums = [str(int(n)) if n.is_integer() else str(n) for n in vals]
        return "X".join(clean_nums)
        
    # 2. Fallback for single numbers (like "70CUT" or "48")
    numbers = re.findall(r'\d+(?:\.\d+)?', s)
    if numbers:
        n = float(numbers[0])
        return str(int(n)) if n.is_integer() else str(n)
        
    return s.replace(' ', '')

def get_clean_col_name(col):
    return str(col).strip().upper().replace(' ', '')

def find_col_exact(df, possible_names):
    possible_names_clean = [n.replace(' ', '') for n in possible_names]
    for col in df.columns:
        if get_clean_col_name(col) in possible_names_clean:
            return col
    return None

# --- Auto-Sync Engine (FLEXIBLE PO + NORMALIZED SIZE MAPPING) ---
def get_extracted_po_dictionary(history_ids):
    if not history_ids: return {}
    
    conn = sqlite3.connect(DB_NAME)
    placeholders = ','.join('?' * len(history_ids))
    query = f"SELECT po_number, size, quantity FROM extracted_data WHERE history_id IN ({placeholders})"
    
    df_extracted = pd.read_sql_query(query, conn, params=tuple(history_ids))
    conn.close()
    
    if df_extracted.empty: return {}
    
    # Clean data and mathematically normalize all sizes!
    df_extracted['quantity'] = pd.to_numeric(df_extracted['quantity'].astype(str).str.replace(',', ''), errors='coerce')
    df_extracted['po_number'] = df_extracted['po_number'].astype(str).str.strip().str.upper()
    df_extracted['size'] = df_extracted['size'].apply(normalize_size)
    
    df_extracted = df_extracted.dropna(subset=['quantity'])
    
    df_summed = df_extracted.groupby(['po_number', 'size'])['quantity'].sum().reset_index()
    
    summed_dict = {}
    for _, row in df_summed.iterrows():
        summed_dict[(row['po_number'], row['size'])] = row['quantity']
        
    return summed_dict

def sync_extracted_data_to_master(df_master, summed_dict):
    updates_made = 0
    if not summed_dict: return df_master, 0
    
    po_col = find_col_exact(df_master, ["PURCHASEORDER", "PO", "PONUMBER", "PONO.", "PONO"])
    size_col = find_col_exact(df_master, ["SIZE", "DIMENSION", "DIMENSIONS"])
    del_rem_col = find_col_exact(df_master, ["DELIVERREMAINDER", "DELIVERYREMAINDER"])
    out_col_name = find_col_exact(df_master, ["OUTSTANDING"]) or 'OUTSTANDING'
    
    if '#inv' in df_master.columns: df_master['#inv'] = 0.0
    if del_rem_col and del_rem_col in df_master.columns:
        df_master[del_rem_col] = pd.to_numeric(df_master[del_rem_col].astype(str).str.replace(',', ''), errors='coerce').fillna(0.0)
    if out_col_name in df_master.columns and del_rem_col: 
        df_master[out_col_name] = df_master[del_rem_col]
        
    remaining_qty_pool = {}
    if not size_col:
        for (p, s), qty in summed_dict.items():
            remaining_qty_pool[(p, "")] = remaining_qty_pool.get((p, ""), 0.0) + qty
    else:
        remaining_qty_pool = summed_dict.copy()
    
    po_size_counts = {}
    po_only_counts = {}
    if po_col:
        for index, row in df_master.iterrows():
            po = str(row[po_col]).strip().upper()
            size_raw = row[size_col] if size_col else ""
            size_val = normalize_size(size_raw)
            delivery_rem = float(row[del_rem_col]) if del_rem_col else 0.0
            
            if delivery_rem != 0:
                po_size_counts[(po, size_val)] = po_size_counts.get((po, size_val), 0) + 1
                po_only_counts[po] = po_only_counts.get(po, 0) + 1
    
    if po_col:
        for index, row in df_master.iterrows():
            po = str(row[po_col]).strip().upper()
            size_raw = row[size_col] if size_col else ""
            size_val = normalize_size(size_raw)
            delivery_rem = float(row[del_rem_col]) if del_rem_col else 0.0
            
            fill_amount = 0.0
            updated = False
            
            if delivery_rem != 0:
                if po in po_only_counts: po_only_counts[po] -= 1
                if (po, size_val) in po_size_counts: po_size_counts[(po, size_val)] -= 1
            
            # Exact Match
            if (po, size_val) in remaining_qty_pool and remaining_qty_pool[(po, size_val)] > 0:
                available_qty = remaining_qty_pool[(po, size_val)]
                if po_size_counts.get((po, size_val), 0) <= 0: fill_amount = available_qty
                else: fill_amount = min(available_qty, delivery_rem) if delivery_rem > 0 else 0
                remaining_qty_pool[(po, size_val)] -= fill_amount
                updated = True
                
            # Fallback for invoices with no size
            elif (po, "") in remaining_qty_pool and remaining_qty_pool[(po, "")] > 0:
                available_qty = remaining_qty_pool[(po, "")]
                if po_only_counts.get(po, 0) <= 0: fill_amount = available_qty
                else: fill_amount = min(available_qty, delivery_rem) if delivery_rem > 0 else 0
                remaining_qty_pool[(po, "")] -= fill_amount
                updated = True

            if updated:
                df_master.at[index, '#inv'] = fill_amount
                df_master.at[index, out_col_name] = delivery_rem - fill_amount
                updates_made += 1
            else:
                df_master.at[index, '#inv'] = 0.0
                df_master.at[index, out_col_name] = delivery_rem

    # Formatting
    if '#inv' in df_master.columns:
        df_master['#inv'] = df_master['#inv'].astype(float).map('{:.2f}'.format)
    if out_col_name in df_master.columns:
        df_master[out_col_name] = df_master[out_col_name].astype(float).map('{:.2f}'.format)
    if del_rem_col and del_rem_col in df_master.columns:
        df_master[del_rem_col] = df_master[del_rem_col].astype(float).map('{:.2f}'.format)
                
    return df_master, updates_made

def update_excel_template_in_memory(raw_file_bytes, summed_dict):
    wb = openpyxl.load_workbook(io.BytesIO(raw_file_bytes))
    ws = wb.active 
    
    header_row = 1
    col_map = {}
    for r in range(1, 15):
        row_vals = [ws.cell(row=r, column=c).value for c in range(1, ws.max_column + 1)]
        row_vals_clean = [get_clean_col_name(v) for v in row_vals]
        
        if "PURCHASEORDER" in row_vals_clean or "PO" in row_vals_clean or "PONO." in row_vals_clean:
            header_row = r
            for c, val in enumerate(row_vals, start=1):
                if isinstance(val, str): col_map[get_clean_col_name(val)] = c
            break

    def get_col_idx(possible_names):
        possible_names_clean = [n.replace(' ', '') for n in possible_names]
        for name in possible_names_clean:
            if name in col_map: return col_map[name]
        return None
            
    po_col = get_col_idx(["PURCHASEORDER", "PO", "PONUMBER", "PONO.", "PONO"])
    size_col = get_col_idx(["SIZE", "DIMENSION", "DIMENSIONS"])
    inv_col = get_col_idx(["#INV", "INV", "INVOICEQTY"])
    del_rem_col = get_col_idx(["DELIVERREMAINDER", "DELIVERYREMAINDER"])
    out_col = get_col_idx(["OUTSTANDING"])

    remaining_qty_pool = {}
    if not size_col:
        for (p, s), qty in summed_dict.items():
            remaining_qty_pool[(p, "")] = remaining_qty_pool.get((p, ""), 0.0) + qty
    else:
        remaining_qty_pool = summed_dict.copy()

    po_size_counts = {}
    po_only_counts = {}
    if po_col:
        for r in range(header_row + 1, ws.max_row + 1):
            po_raw = ws.cell(row=r, column=po_col).value
            po_val = "" if po_raw is None else str(po_raw).strip().upper()
            
            if po_val and po_val != 'NONE':
                size_raw = ws.cell(row=r, column=size_col).value if size_col else ""
                size_val = normalize_size(size_raw)
                
                del_rem_val = 0.0
                if del_rem_col:
                    raw_rem = ws.cell(row=r, column=del_rem_col).value
                    try:
                        if isinstance(raw_rem, str): del_rem_val = float(raw_rem.replace(',', ''))
                        elif raw_rem is not None: del_rem_val = float(raw_rem)
                    except: pass
                    
                if del_rem_val != 0:
                    po_size_counts[(po_val, size_val)] = po_size_counts.get((po_val, size_val), 0) + 1
                    po_only_counts[po_val] = po_only_counts.get(po_val, 0) + 1
    
    if po_col:
        for r in range(header_row + 1, ws.max_row + 1):
            po_raw = ws.cell(row=r, column=po_col).value
            po_val = "" if po_raw is None else str(po_raw).strip().upper()
            size_raw = ws.cell(row=r, column=size_col).value if size_col else ""
            size_val = normalize_size(size_raw)
            
            if inv_col: 
                c_inv = ws.cell(row=r, column=inv_col)
                c_inv.value = 0.0
                c_inv.number_format = '0.00'
            
            del_rem_val = 0.0
            if del_rem_col:
                raw_rem = ws.cell(row=r, column=del_rem_col).value
                try:
                    if isinstance(raw_rem, str): del_rem_val = float(raw_rem.replace(',', ''))
                    elif raw_rem is not None: del_rem_val = float(raw_rem)
                except: del_rem_val = 0.0
                
                c_del = ws.cell(row=r, column=del_rem_col)
                c_del.number_format = '0.00'
                
            if out_col: 
                c_out = ws.cell(row=r, column=out_col)
                c_out.value = del_rem_val
                c_out.number_format = '0.00'
            
            fill_amount = 0.0
            updated = False
            
            if del_rem_val != 0:
                if po_val in po_only_counts: po_only_counts[po_val] -= 1
                if (po_val, size_val) in po_size_counts: po_size_counts[(po_val, size_val)] -= 1
            
            if (po_val, size_val) in remaining_qty_pool and remaining_qty_pool[(po_val, size_val)] > 0:
                available_qty = remaining_qty_pool[(po_val, size_val)]
                if po_size_counts.get((po_val, size_val), 0) <= 0: fill_amount = available_qty
                else: fill_amount = min(available_qty, del_rem_val) if del_rem_val > 0 else 0
                remaining_qty_pool[(po_val, size_val)] -= fill_amount
                updated = True
                
            elif (po_val, "") in remaining_qty_pool and remaining_qty_pool[(po_val, "")] > 0:
                available_qty = remaining_qty_pool[(po_val, "")]
                if po_only_counts.get(po_val, 0) <= 0: fill_amount = available_qty
                else: fill_amount = min(available_qty, del_rem_val) if del_rem_val > 0 else 0
                remaining_qty_pool[(po_val, "")] -= fill_amount
                updated = True
            
            if updated:
                if inv_col: 
                    c_inv = ws.cell(row=r, column=inv_col)
                    c_inv.value = float(fill_amount)
                    c_inv.number_format = '0.00'
                if out_col: 
                    c_out = ws.cell(row=r, column=out_col)
                    c_out.value = float(del_rem_val - fill_amount)
                    c_out.number_format = '0.00'

    owner_col = col_map.get("OWNER")
    if owner_col: ws.delete_cols(owner_col)

    output_buffer = io.BytesIO()
    wb.save(output_buffer)
    return output_buffer.getvalue()

init_db()

# --- REGEX PATTERNS & EXTRACTION LOGIC ---
po_pattern = re.compile(r'(?<![A-Za-z])(?:PO\s*#|PO\s*NO\.?|PO\s*NUMBER|P\.O\.|P\.O\s*:|ORDER NO\.?:?)\s*[:\-,]?\s*([A-Za-z0-9\-_]+)|(KPO[A-Za-z0-9]+|KPS[A-Za-z0-9]+)', re.IGNORECASE)
qty_pattern = re.compile(r'([\d,]+(?:\.\d+)?)\s*(YDS|YARDS|PCS|YARD|PC)\b', re.IGNORECASE)
size_pattern = re.compile(r'(\d+(?:\.\d+)?\s*[a-zA-Z]*\s*[xX*]\s*\d+(?:\.\d+)?\s*[a-zA-Z]*)', re.IGNORECASE)

def extract_from_text(text, filename):
    text = text.replace('\r', '\n')
    lines = text.split('\n')
    current_po = "Unknown PO"
    extracted_pairs = []
    has_packing_list = any("PACKING LIST" in l.upper() for l in lines)
    parsing_active = not has_packing_list 
    
    for i, line in enumerate(lines):
        line_upper = line.upper()
        if has_packing_list:
            if "PACKING LIST" in line_upper: parsing_active = True
            elif "COMMERCIAL INVOICE" in line_upper: parsing_active = False
        if not parsing_active: continue
        
        po_match = po_pattern.search(line)
        if po_match and "ITEM" not in line_upper:
            current_po = po_match.group(1) if po_match.group(1) else po_match.group(2)
            current_po = current_po.strip().upper()
            
        if any(x in line_upper for x in ["TOTAL", "SUBTOT", "==="]): continue
        if "F.O.C" in line_upper or "FOC" in line_upper: continue
        
        current_size = ""
        size_match = size_pattern.search(line_upper)
        if size_match: current_size = normalize_size(size_match.group(1))
            
        for q in qty_pattern.findall(line):
            if '.' not in str(q[0]): 
                extracted_pairs.append({"File Name": filename, "PO number": current_po, "Size": current_size, "QTY": f"{q[0]}"})
    return extracted_pairs

def process_tabular_data(df, filename):
    grid = df.fillna("").astype(str).values.tolist()
    extracted_pairs = []
    current_po = "Unknown PO"
    qty_col_indices = set()
    po_col_index = -1
    size_col_index = -1
    
    for row in grid:
        clean_row = [str(cell).strip() for cell in row]
        row_str = " ".join(clean_row).upper()
        row_has_header = False
        
        for i, cell in enumerate(clean_row):
            cell_upper = cell.upper().replace(' ', '')
            
            if cell_upper in ["YARD", "YARDS", "YDS", "PCS", "PC", "QUANTITY", "QTY", "PURCHASEQTY", "Q'TY"]:
                qty_col_indices.add(i)
                if "COLOR" in row_str or "ITEM" in row_str or "SIZE" in row_str or "STYLE" in row_str or cell_upper in ["QUANTITY", "QTY", "PURCHASEQTY", "Q'TY"]:
                    row_has_header = True
            
            if cell_upper in ["PO", "P.O.", "PO#", "PONO.", "PONO", "PURCHASEORDER", "ORDERNO.", "ORDERNO", "PONUMBER"]:
                po_col_index = i
                row_has_header = True
                
            if cell_upper in ["SIZE", "DIMENSION", "DIMENSIONS"]:
                size_col_index = i
                row_has_header = True
                
        if row_has_header: continue
        if any(x in row_str for x in ["TOTAL", "SUBTOT", "F.O.C", "FOC", "==="]): continue
        
        po_found = False
        if po_col_index != -1 and po_col_index < len(clean_row) and clean_row[po_col_index]:
            current_po = clean_row[po_col_index].upper()
            po_found = True
            
        if not po_found:
            po_match = po_pattern.search(row_str)
            if po_match:
                current_po = po_match.group(1) if po_match.group(1) else po_match.group(2)
                current_po = current_po.upper()
            else:
                for i, cell in enumerate(clean_row):
                    if cell.upper().replace(' ', '') in ["PO", "P.O.", "PO#", "PONO.", "PURCHASEORDER"]:
                        if i + 1 < len(clean_row) and clean_row[i+1]: 
                            current_po = clean_row[i+1].upper()
                            break
                            
        current_size = ""
        if size_col_index != -1 and size_col_index < len(clean_row):
            current_size = normalize_size(clean_row[size_col_index])
                            
        extracted_from_grid = False
        for col_idx in qty_col_indices:
            if col_idx < len(clean_row) and clean_row[col_idx]:
                clean_num = clean_row[col_idx].replace(',', '')
                try:
                    if '.' not in clean_num and float(clean_num) > 0:
                        extracted_pairs.append({"File Name": filename, "PO number": current_po, "Size": current_size, "QTY": clean_row[col_idx]})
                        extracted_from_grid = True
                except: pass
                
        if not extracted_from_grid:
            for cell in clean_row:
                for q in qty_pattern.findall(cell):
                    if '.' not in str(q[0]):
                        extracted_pairs.append({"File Name": filename, "PO number": current_po, "Size": current_size, "QTY": q[0]})
                        
    return extracted_pairs

# --- UI START ---
PAGE_EXTRACT = ":material/plumbing: Extraction Tool"
PAGE_HISTORY = ":material/database: Extraction Database"
PAGE_MASTER = ":material/account_tree: Master PO Line Database"

st.sidebar.title(":material/menu: Navigation")
page = st.sidebar.radio("Select a Module:", [PAGE_EXTRACT, PAGE_HISTORY, PAGE_MASTER])

st.sidebar.markdown("---")
st.sidebar.subheader(":material/build: Advanced")
if st.sidebar.button(":material/delete_forever: Reset Database", type="primary", help="Clears old extraction history."):
    try: 
        os.remove(DB_NAME)
        st.sidebar.success("Database erased! Please refresh page.")
    except: 
        st.sidebar.error("Could not delete. Close other files and try again.")

if page == PAGE_EXTRACT:
    st.title(":material/plumbing: Invoice & Packing List Extractor")
    st.markdown("Upload files to extract **PO numbers**, **Sizes**, and **Quantities**.")
    
    with st.container(border=True):
        uploaded_files = st.file_uploader("Upload Document Files", accept_multiple_files=True, type=['pdf', 'docx', 'doc', 'csv', 'xlsx'])
        if st.button(":material/play_arrow: Analyze Files", type="primary"):
            if not uploaded_files: st.warning("Please upload at least one file.")
            else:
                all_results = []
                with st.spinner('Processing files with Data Engines...'):
                    for file in uploaded_files:
                        filename = file.name
                        file_results = []
                        try:
                            if filename.lower().endswith(('.csv', '.xlsx')):
                                if filename.lower().endswith('.csv'):
                                    try: df = pd.read_csv(file, header=None, on_bad_lines='skip')
                                    except: 
                                        file.seek(0)
                                        df = pd.read_csv(file, header=None, encoding='latin1', on_bad_lines='skip')
                                else: df = pd.read_excel(file, header=None)
                                file_results = process_tabular_data(df, filename)
                            else:
                                text = ""
                                if filename.lower().endswith('.pdf'):
                                    reader = PdfReader(file)
                                    text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                                elif filename.lower().endswith('.docx'):
                                    doc = docx.Document(file)
                                    text = "\n".join([para.text for para in doc.paragraphs])
                                    for table in doc.tables:
                                        for row in table.rows: text += "\n" + " ".join([cell.text for cell in row.cells])
                                elif filename.lower().endswith('.doc'):
                                    text = file.getvalue().decode('ascii', errors='ignore')
                                file_results = extract_from_text(text, filename)
                                
                            if file_results:
                                file_results_df = pd.DataFrame(file_results).drop_duplicates()
                                log_to_db(filename, file_results_df)
                                all_results.extend(file_results_df.to_dict('records'))
                        except Exception as e: st.error(f"Error processing {filename}: {e}")

                if all_results:
                    df_results = pd.DataFrame(all_results).drop_duplicates()
                    st.success(f"Extraction Complete! Found {len(df_results)} pairs.")
                    st.dataframe(df_results, use_container_width=True, hide_index=True)
                    buffer = io.BytesIO()
                    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                        df_results.to_excel(writer, index=False, sheet_name='Extracted Data')
                    st.download_button(label=":material/download: Download This Batch as Excel", data=buffer.getvalue(), file_name='paired_po_quantities.xlsx', mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                else: st.info("Analysis complete, but no matching PO Numbers or Quantities were found.")

elif page == PAGE_HISTORY:
    st.title(":material/database: Extraction History")
    df = fetch_db_report()
    if df.empty: st.info("No history found in database.")
    else:
        with st.container(border=True):
            history_df = df.copy()
            history_df.insert(0, "Delete", False)
            edited_df = st.data_editor(history_df, column_config={"Delete": st.column_config.CheckboxColumn("Select to Delete")}, hide_index=True, use_container_width=True)
            if st.button(":material/delete: Delete Selected"):
                to_delete = edited_df[edited_df['Delete'] == True]
                if not to_delete.empty:
                    for index, row in to_delete.iterrows(): delete_history_record(row['id'])
                    st.rerun()

elif page == PAGE_MASTER:
    st.title(":material/account_tree: Master PO Line Database")
    
    with st.container(border=True):
        st.subheader(":material/upload_file: Upload Master Report")
        po_line_file = st.file_uploader("Upload Excel/CSV Tracking Report", type=['csv', 'xlsx'], key="po_up")
        if po_line_file:
            try:
                with st.spinner("Saving EXACT file to Database..."):
                    if po_line_file.name.lower().endswith('.csv'):
                        try: df_master = pd.read_csv(po_line_file)
                        except: 
                            po_line_file.seek(0)
                            df_master = pd.read_csv(po_line_file, encoding='latin1')
                    else: df_master = pd.read_excel(po_line_file)
                    
                    df_master.columns = [str(col).strip() for col in df_master.columns]
                    
                    if 'QTY' in df_master.columns and 'File Name' in df_master.columns:
                        st.error("Oops! You accidentally uploaded an Extracted Data file here. Upload tracking reports only.")
                        st.stop()

                    upload_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    df_master['Source_File_Name'] = po_line_file.name
                    df_master['Upload_Date'] = upload_date

                    conn = sqlite3.connect(DB_NAME)
                    cursor = conn.cursor()
                    try: cursor.execute("DELETE FROM master_po_data WHERE Source_File_Name = ?", (po_line_file.name,))
                    except: pass
                    cursor.execute("DELETE FROM master_files WHERE file_name = ?", (po_line_file.name,))
                    
                    raw_file_bytes = po_line_file.getvalue()
                    cursor.execute("INSERT INTO master_files (file_name, upload_date, file_data) VALUES (?, ?, ?)", (po_line_file.name, upload_date, raw_file_bytes))
                    df_master.to_sql('master_po_data', conn, if_exists='append', index=False)
                    conn.commit()
                    conn.close()
                st.success(f"Saved '{po_line_file.name}' to the database.")
            except Exception as e: st.error(f"System Error: {e}")

    st.markdown("---")
    st.subheader(":material/manage_accounts: File Mapping & Auto-Sync")
    summary_df = get_master_files_summary()
    
    if summary_df.empty: 
        st.info("No master files have been uploaded yet.")
    else:
        st.dataframe(summary_df, use_container_width=True, hide_index=True)
        
        st.markdown("#### :material/account_tree: Step 1: Select Files to Map")
        col_m1, col_m2 = st.columns(2)
        
        with col_m1:
            file_to_manage = st.selectbox("1. Select Master File to Update:", summary_df['File Name'].tolist())
            
        with col_m2:
            history_df = fetch_db_report()
            if not history_df.empty:
                history_options = history_df.apply(lambda row: f"[ID: {row['id']}] {row['File Name']} ({row['PO Rows Extracted']} POs)", axis=1).tolist()
                selected_histories = st.multiselect("2. Select Extraction Batch(es) to Inject:", history_options, default=[history_options[0]])
                selected_history_ids = [int(s.split("]")[0].split(":")[1].strip()) for s in selected_histories]
            else:
                st.warning("No extracted data found. Go to 'Extraction Tool' first.")
                selected_history_ids = []
                
        st.markdown("#### :material/play_circle: Step 2: Execute Operations")
        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            if st.button(":material/visibility: Preview Data", use_container_width=True):
                st.session_state.preview_active = True
                st.session_state.sync_active = False
                st.session_state.preview_file = file_to_manage
        with col2:
            if st.button(":material/sync: Auto-Sync Extracted POs", use_container_width=True, type="primary"):
                st.session_state.sync_active = True
                st.session_state.preview_active = False 
                st.session_state.sync_file = file_to_manage
        with col3:
            if st.button(":material/delete: Delete Master File", use_container_width=True):
                delete_master_file(file_to_manage)
                st.session_state.preview_active = False 
                st.session_state.sync_active = False
                st.rerun() 
                
        if st.session_state.get('preview_active') and st.session_state.get('preview_file') == file_to_manage:
            with st.container(border=True):
                st.markdown(f"**Previewing:** `{file_to_manage}`")
                file_data_df = preview_master_file(file_to_manage)
                st.data_editor(file_data_df, use_container_width=True, hide_index=True, disabled=True)
                raw_file_data = get_raw_master_file(file_to_manage)
                if raw_file_data:
                    mime_type = 'text/csv' if file_to_manage.lower().endswith('.csv') else 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                    st.download_button(label=":material/download: Download Exact Original File", data=raw_file_data, file_name=file_to_manage, mime=mime_type)
        
        if st.session_state.get('sync_active') and st.session_state.get('sync_file') == file_to_manage:
            with st.container(border=True):
                st.markdown(f"### :material/task_alt: Auto-Synced Data (Smart Exact Match)")
                
                if not selected_history_ids:
                    st.error("Please select at least one Extraction Batch from 'Step 1' above before syncing.")
                else:
                    file_data_df = preview_master_file(file_to_manage)
                    summed_dict = get_extracted_po_dictionary(selected_history_ids)
                    
                    synced_df, updates_count = sync_extracted_data_to_master(file_data_df, summed_dict)
                    
                    if updates_count > 0: st.success(f"Successfully mapped and calculated {updates_count} exact PO lines!")
                    else: st.warning("No matching PO numbers / Sizes were found in the selected extraction batches.")
                        
                    st.data_editor(synced_df, use_container_width=True, hide_index=True, disabled=True)
                    
                    raw_file_bytes = get_raw_master_file(file_to_manage)
                    if raw_file_bytes and file_to_manage.lower().endswith('.xlsx'):
                        updated_excel_bytes = update_excel_template_in_memory(raw_file_bytes, summed_dict)
                        st.download_button(label=":material/download: Download Perfect Template (.xlsx)", data=updated_excel_bytes, file_name=f"UPDATED_{file_to_manage}", mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', type="primary")
                    else:
                        download_df = synced_df.drop(columns=['Source_File_Name', 'Upload_Date', 'Owner'], errors='ignore')
                        csv_buffer = download_df.to_csv(index=False).encode('utf-8')
                        st.download_button(label=":material/download: Download Updated File (.csv)", data=csv_buffer, file_name=f"UPDATED_{file_to_manage}", mime='text/csv', type="primary")